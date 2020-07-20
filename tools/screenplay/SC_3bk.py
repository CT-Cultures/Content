# -*- coding: utf-8 -*-
"""
Created on Mon Feb 17 00:31:27 2020

@author: VXhpUS
"""
import numpy as np
import pandas as pd
import re
from docx import Document
import time
from xml.dom import minidom
import lxml.etree as et
import codecs

#%%
class Screenplay(object):
    
    def __init__(self, sc: str = None):      
        super(Screenplay, self).__init__()
        self.sc = sc
        self.elements = Elements()
        self.reformat = Reformat()
        
        # Scene heading element
    
    def open(self, scpath: str):

        pass
    
    def read_docx(self, filepath: str, sctype: str =None) -> pd.DataFrame:
        '''
        This function reads word docx into a screenplay structured pd.DataFrame, 
        oneline per row.

        Parameters
        ----------
        filepath : str
            DESCRIPTION.
        sctype : str, optional
            DESCRIPTION. The default is None.

        Returns
        -------
        script : pd.DataFrame 
            DESCRIPTION.

        '''
        document = Document(filepath)
        num_p = len(document.paragraphs)
        script = []
        for num in range(num_p):
            script.append(document.paragraphs[num].text)
        script = pd.Series(script).rename('pcontent')
        script = script[script != ''] # remove empty lines
        script = pd.DataFrame(script) # convert to DataFrame
        
        # Paragraph type
        '''
        h = scene heading
            - Setting (i.e. INT/EXT)
            - Parenthesis (i.e. (Flashback))
        d = dialog
           - Characters
           - Parenthesis
        t = text
        s = shot (i.e. POV)
        t = transition (i.e. CUT TO, DISSOLVE, FADE, CROSS FADE)
            
        '''
        script['ptype'] = np.nan 
        
        # Scene Heading elements
        script['h_number'] = None
        script['h_inout'] = None
        script['h_title'] = None
        script['h_time'] = None
        script['h_parenthesis'] = None
        script['h_characters_in_scene'] = None
        
        # Dialog 
        script['d_character'] = None
        script['d_character_parenthesis'] = None
        script['d_dialog'] = None
        script['d_dialog_parenthesis'] = None
        return script
    
    def export_csv(self, filepath):
        self.sc.to_csv(filepath, encoding='utf-8-sig')
    
    
class Elements(object):
    
    def __init__(self):
        super(Elements, self).__init__()
        
    def identify_scene_heading(self, sc, pattern: str = '\d*[.\s]') -> pd.DataFrame:
        sc.loc[sc['pcontent'].str.contains(pattern), 'ptype'] = 'h'
        return sc
    
    def identify_dialog(self, sc, pattern: str = u'[：:]') -> pd.DataFrame:
        sc.loc[sc['pcontent'].apply(lambda x: x[0:10]).str.contains(pattern), 'ptype'] = 'd'
        return sc
    
    def identify_action(self, sc, pattern: str = None) -> pd.DataFrame:
        if pattern:
            sc.loc[sc['pcontent'].str.contains(pattern), 'ptype'] = 'a'
        else:
            sc['ptype'].fillna(value='a', inplace=True)
        return sc
    
    # Scene Heading Subelements
    def identify_scene_number(self, 
                              sc, 
                              pattern: str, 
                              identify_type_first: bool = False) -> pd.DataFrame:
        if identify_type_first:
            sc = self.identify_scene_heading(sc, pattern)
        sc.loc[sc['ptype'] == 'h', 'h_number'] = sc.loc[sc['ptype'] == 'h', 'pcontent'].str.extract(pattern).iloc[:, 0]
        #sc['h_number'].fillna(method='ffill', inplace=True)
        return sc
    
    def identify_scene_inout(self, 
                              sc, 
                              pattern: str = r"[\w\S]*(.*?)[，,]", 
                              identify_type_first: bool = False) -> pd.DataFrame:
        if identify_type_first:
            sc = self.identify_scene_heading(sc, pattern)
        sc.loc[sc['ptype'] == 'h', 'h_inout'] = sc.loc[sc['ptype'] == 'h', 'pcontent'].str.extract(pattern).iloc[:, 0]
        #sc['h_inout'].fillna(method='ffill', inplace=True)
        return sc
        pass
        
    def identify_scene_title(self, 
                          sc, 
                          pattern: str = r'景[，,](\s*\w*.*)[，,]', 
                          identify_type_first: bool = False) -> pd.DataFrame:
        if identify_type_first:
            sc = self.identify_scene_heading(sc, pattern)
        sc.loc[sc['ptype'] == 'h', 'h_title'] = sc.loc[sc['ptype'] == 'h', 'pcontent'].str.extract(pattern).iloc[:, 0]
        #sc['h_title'].fillna(method='ffill', inplace=True)
        return sc

    def identify_scene_time(self, 
                          sc, 
                          pattern: str = r'[，,](\s*\w*.*$)', 
                          identify_type_first: bool = False) -> pd.DataFrame:
        if identify_type_first:
            sc = self.identify_scene_heading(sc, pattern)
        sc.loc[sc['ptype'] == 'h', 'h_time'] = sc.loc[sc['ptype'] == 'h', 'pcontent'].str.extract(pattern).iloc[:, 0]
        #sc['h_time'].fillna(method='ffill', inplace=True)
        return sc
    
    # Dialog Subelements
    def identify_dialog_character(self,
                                  sc,
                                  pattern: str = r'(^.*$)[:：]',
                                  identify_type_first: bool = False) -> pd.DataFrame:
        if identify_type_first:
            sc = self.identify_dialog(sc, pattern)
        sc.loc[sc['ptype'] == 'd', 'd_character'] = sc.loc[sc['ptype'] == 'd', 'pcontent'].str.extract(pattern).iloc[:, 0]
        sc  = self.identify_dialog_character_parenthesis(sc)
        #sc['h_time'].fillna(method='ffill', inplace=True)
        return sc

    def identify_dialog_character_parenthesis(self,
                                  sc: pd.DataFrame,
                                  pattern: str = r'[（(](.*)[）)]',
                                  identify_type_first: bool = False) -> pd.DataFrame:
        if identify_type_first:
            sc = self.identify_dialog(sc, pattern)
        sc.loc[sc['ptype'] == 'd', 'd_character_parenthesis'] = sc.loc[sc['ptype'] == 'd', 'd_character'].str.extract(pattern).iloc[:, 0]
        sc.loc[sc['ptype'] == 'd', 'd_character'] = sc.loc[sc['ptype'] == 'd', 'd_character'].str.replace(pattern, '')
        #sc['h_time'].fillna(method='ffill', inplace=True)
        return sc        
                        
    def identify_dialog_dialog(self,
                                 sc,
                                 pattern: str = r'[:：](.*)',
                                 identify_type_first: bool = False) -> pd.DataFrame:
       if identify_type_first:
           sc = self.identify_dialog(sc, pattern)
       sc.loc[sc['ptype'] == 'd', 'd_dialog'] = sc.loc[sc['ptype'] == 'd', 'pcontent'].str.extract(pattern).iloc[:, 0]
       sc.loc[sc['ptype'] == 'd', 'd_dialog'] = sc.loc[sc['ptype'] == 'd', 'd_dialog'].apply(lambda x: x.lstrip(' ').rstrip(' '))
       sc  = self.identify_dialog_parenthetisis(sc)
       return sc
   
    def identify_dialog_parenthetisis(self,
                                      sc: pd.DataFrame,
                                      pattern: str = r'[（(](.*)[）)]',
                                      identify_type_first: bool = False) -> pd.DataFrame:
       if identify_type_first:
           sc = self.identify_dialog(sc, pattern)
       sc.loc[sc['ptype'] == 'd', 'd_dialog_parenthesis'] = sc.loc[sc['ptype'] == 'd', 'd_dialog']. \
           str.extractall(r'[（(](.*?)[）)]').reset_index().groupby('level_0')[0].agg(';'.join).str.split(';')


       #sc['h_time'].fillna(method='ffill', inplace=True)
       return sc
   
######################################################
class Reformat(object):
    
    def __init__(self):
        super(Reformat, self).__init__()      
    
    def to_openxml(self, sc: pd.DataFrame, save: bool = False, file_path: str  = None) -> str:
        '''
        

        Parameters
        ----------
        sc : pd.DataFrame
            DESCRIPTION.
        save : str, optional
            DESCRIPTION. The default is None.
        file_path : str, optional
            DESCRIPTION. The default is None.

        Returns
        -------
        str
            DESCRIPTION.

        '''
        sc['formatted'] = None
        element2basestyle = {
                            'text':'Normal Text',
                            'h': 'Scene Heading',
                            'a': 'Action',
                            'c': 'Character',
                            'p': 'Parenthetical',
                            'd': 'Dialogue',
                            't': 'Transition',
                            's': 'Shot',
                            'uc': 'Unspoken Character'
                            }
        header_xml = '<?xml version="1.0" encoding="UTF-8"?>\n'
        header_document = '<document version="30" type="Open Screenplay Format document">\n' \
            +   ' <info pagecount="" uuid=""8927aa47-d7f7-4110-9d7d-f385abd8e001""/>\n' \
            +   ' <settings show_all_revisions="false" show_revisions="false" revision="0" dialogue_locked="false"' \
            +   ' dialoguenumber_format="(#)" dialoguenumber_start="1" dialogue_numbering="false" scenes_locked="false"' \
            +   ' scenenumber_position="3" scenenumber_format="#" scenenumber_start="1" scenenumber_skip_io="false"' \
            +   ' scenenumber_mode="1AB" scene_numbering="false" pagenumber_mode="1AB" pagenumber_start="1"' \
            +   ' pages_locked="false" footer_first_page="false" header_first_page="false" footer_alignment="3"' \
            +   ' header_alignment="3" page_footer="" page_header="#." scene_time_separator=" - " omitted_text="OMITTED"' \
            +   ' auto_omit_scenes="false" number_continued="true" continued_text="CONTINUED" scenes_continue="false"' \
            +   ' more_text="(MORE)" cont_text="(cont\'d)" dialogue_pagebreaks="true" dialogue_continues="true"' \
            +   'break_on_sentences="true" element_spacing="1.00" normal_linesperinch="6.0" margin_right="317"' \
            +   'margin_left="317" margin_bottom="220" margin_top="317" page_height="2794" page_width="2159"/>' \
            +   ' <fadein_settings navigator_pagecount="2" navigator_show="1" index_cards_use_colors="true"' \
            +   'index_cards_use_folders="true" index_cards_text_size="1" index_cards_show="2" last_position="9, 0"' \
            +   'saved_with=""/>\n'
            
        header_format = ' <styles>\n' \
            +   '  <style size="12" font="Courier New" label="Normal Text" builtin_index="0" builtin="1" name="Normal Text"/>\n' \
            +   '  <style size="12" font="Courier New" label="Scene Heading" builtin_index="1" builtin="1" name="Scene Heading"' \
            +   'allcaps="1" keepwithnext="1" spacebefore="2.0" style_tab_after="Action" style_enter="Action" basestylename="Normal Text"/>' \
            +   '  <style size="12" font="Courier New" label="Action" builtin_index="2" builtin="1" name="Action" spacebefore="1.0"' \
            +   'basestylename="Normal Text" style_tab_before="Character"/>\n' \
            +   '  <style size="12" font="Courier New" label="Character" builtin_index="3" builtin="1" name="Character" allcaps="1"' \
            +   'keepwithnext="1" spacebefore="1.0" style_tab_after="Parenthetical" style_enter="Dialogue" basestylename="Normal Text"' \
            +   'style_tab_before="Action" leftindent="635"/>\n' \
            +   '  <style size="12" font="Courier New" label="Parenthetical" builtin_index="4" builtin="1" name="Parenthetical" keepwithnext="1"' \
            +   'style_tab_after="Dialogue" style_enter="Dialogue" basestylename="Normal Text" style_tab_before="Dialogue" leftindent="508"' \
            +    ' rightindent="508"/>\n' \
            +   '  <style size="12" font="Courier New" label="Dialogue" builtin_index="5" builtin="1" name="Dialogue"' \
            +   'style_tab_after="Parenthetical" style_enter="Action" basestylename="Normal Text" style_tab_before="Parenthetical"' \
            +   'leftindent="330" rightindent="254"/>\n' \
            +   '  <style size="12" font="Courier New" label="Transition" builtin_index="6" builtin="1" name="Transition" allcaps="1"' \
            +   'spacebefore="1.0" style_tab_after="Action" style_enter="Scene Heading" basestylename="Normal Text" leftindent="1016"' \
            +   'rightindent="127" align="right"/>\n' \
            +   '  <style size="12" font="Courier New" label="Shot" builtin_index="7" builtin="1" name="Shot" allcaps="1" keepwithnext="1"' \
            +   'spacebefore="1.0" style_tab_after="Action" style_enter="Action" basestylename="Normal Text"/>\n' \
            +   '  <style size="12" font="Courier New" label="Shot" builtin_index="7" builtin="1" name="Shot" allcaps="1" keepwithnext="1"' \
            +   'spacebefore="1.0" style_tab_after="Action" style_enter="Action" basestylename="Normal Text"/>\n' \
            +   '  <header_style basestylename="Normal Text"/>\n' \
            +   '  <footer_style basestylename="Normal Text"/>\n' \
            +   ' </styles>\n'
        header_paragraphs = ' <paragraphs>\n'
        
        footer_paragraphs = ' </paragraphs>\n'
        section_list = ' <list>\n </list>\n'
        footer_document = '</document>\n'
            
        sc = self.heading(sc)
        sc = self.action(sc)
        sc = self.dialog(sc)
        
        formatted_output = header_xml + header_document + header_format + header_paragraphs \
            + sc['formatted'].sum() + footer_paragraphs + section_list + footer_document
            
        if save:
            with codecs.open(file_path, "w", 'utf-8-sig') as f:
                print(formatted_output, file=f)
            f.close()
            
        return formatted_output
    
    
    def heading(self, sc: pd.DataFrame) -> pd.DataFrame:
        sc.loc[sc['ptype'] == 'h', 'formatted'] = \
                sc.loc[sc['ptype'] == 'h', 'h_inout'].apply(lambda x: str(x) + '. ') \
            +   sc.loc[sc['ptype'] == 'h', 'h_title'] \
            +   sc.loc[sc['ptype'] == 'h', 'h_time'].apply(lambda x: ' - ' + str(x))
        
        sc.loc[sc['ptype'] == 'h', 'formatted'] = sc.loc[sc['ptype'] == 'h', 'formatted'].apply(
            lambda x: '  <para>\n   <style basestylename="Scene Heading"/>\n   <text>'
                    + str(x)
                    + '</text>\n  </para>\n'            
            )
        return sc
    
    def action(self, sc: pd.DataFrame) -> pd.DataFrame:
        sc.loc[sc['ptype'] == 'a', 'formatted'] = \
            sc.loc[sc['ptype'] == 'a', 'pcontent'].apply(lambda x: str(x).lstrip('\s').rstrip('\s'))
        sc.loc[sc['ptype'] == 'a', 'formatted'] = sc.loc[sc['ptype'] == 'a', 'formatted'].apply(
            lambda x: '  <para>\n   <style basestylename="Action"/>\n   <text>'
                    + str(x)
                    + '</text>\n  </para>\n'
            )
        return sc

    def dialog(self, sc: pd.DataFrame) -> pd.DataFrame:
        
        #
        sc.loc[sc['ptype'] == 'd', 'formatted'] = \
              sc.loc[sc['ptype'] == 'd', 'd_character'] \
            + sc.loc[sc['ptype'] == 'd', 'd_character_parenthesis'].apply(lambda x: '(' + str(x) + ')' if not pd.isnull(x) else '')
        
        sc.loc[sc['ptype'] == 'd', 'formatted'] = sc.loc[sc['ptype'] == 'd', 'formatted'].apply(
            lambda x: '  <para>\n   <style basestylename="Character"/>\n   <text>'
                    + str(x)
                    + '</text>\n  </para>\n'
            )
             
        sc.loc[(sc['ptype'] == 'd') & (sc['d_dialog_parenthesis'].isna()), 'formatted'] = \
              sc.loc[(sc['ptype'] == 'd') & (sc['d_dialog_parenthesis'].isna()), 'formatted'] \
            + sc.loc[(sc['ptype'] == 'd') & (sc['d_dialog_parenthesis'].isna()), 'd_dialog'].apply(
                lambda x: '  <para>\n   <style basestylename="Parenthetical"/>\n   <text>'
                    + str(x)
                    + '</text>\n  </para>\n'
                )
        
        sc.loc[(sc['ptype'] == 'd') & (~sc['d_dialog_parenthesis'].isna()), 'formatted'] = \
              sc.loc[(sc['ptype'] == 'd') & (~sc['d_dialog_parenthesis'].isna()), 'formatted'] \
            + sc.loc[(sc['ptype'] == 'd') & (~sc['d_dialog_parenthesis'].isna())].agg(format_d_dialog, axis=1)
                        
        return sc
    
    def format_d_dialog(self, x: pd.DataFrame) -> str:
        '''
        formats dialog for XML output, to be used as a function for pandas agg or apply.
        
        Parameters
        ----------
        x : one row or column of pd.DataFrame depending on axis,
            to be used with the pandas agg or apply function
            DESCRIPTION.

        Returns
        -------
        str
            DESCRIPTION.

        '''
        pattern = r'[（(].*?[）)]'
        d_split = x['d_dialog'].split(pattern)
        f_dialog = ''
        for i, d in enumerate(d_split):
            ddialog = '  <para>\n   <style basestylename="Dialogue"/>\n   <text>' \
                + str(d) \
                + '</text>\n  </para>\n'
            if i == 0:
                f_dialog = f_dialog + ddialog
            else:
                dparenthetical = '  <para>\n   <style basestylename="Parenthetical"/>\n   <text>' \
                    + str(x['d_dialog_parenthesis'][i-1]) \
                    + '</text>\n  </para>\n' 
                f_dialog = f_dialog + dparenthetical + ddialog
        return f_dialog


#%% Read and Parse Screenplay
fp = 'SCREENPLAYS_PRIVATE/电影剧本_往来有玉面.docx'
sc = Screenplay()
docx = sc.read_docx(fp)

# Parse Scene Heading
docx = sc.elements.identify_scene_heading(docx, pattern='\#[0-9].*')
docx = sc.elements.identify_dialog(docx)
docx = sc.elements.identify_action(docx)

docx = sc.elements.identify_scene_number(docx, pattern=r'\#(?P<h_number>\d*)')
docx = sc.elements.identify_scene_inout(docx, pattern=r'[0-9*]\d*(\s*\w*.*)景')
docx = sc.elements.identify_scene_title(docx, pattern=r'景[，,](\s*\w*.*)[，,]')
docx = sc.elements.identify_scene_time(docx, pattern=r'[，,].*[，,](\s*\w*.*$)')
                                       
# Parse Dialog
docx = sc.elements.identify_dialog_character(docx, pattern=r'(.*)[:：]')
docx = sc.elements.identify_dialog_dialog(docx, pattern=r'[:：](.*)')
                                       

# Reformat to openxml
output = sc.reformat.to_openxml(docx, save=True, file_path ='OUTPUT/wanglai.xml')
print(output)

#%%
dp = docx.loc[docx['ptype'] == 'd', 'd_dialog'].str.extractall(r'[（(](.*?)[）)]').reset_index().groupby('level_0')[0].agg(';'.join).str.split(';')
dp = dp.rename('d_parenthesis')
docx['d_parenthesis'].update(dp)

d_dialog_w_p = docx.loc[~docx['d_parenthesis'].isna(), 'd_dialog'].str.split(r'[（(].*?[）)]')

d_dialog_w_p
dp

d_dialog = ''
for ds, ps in zip(d_dialog_w_p, dp):
    for i, d in enumerate(ds):
        if i == 0:
            d_dialog = d_dialog + d
        else:
            d_dialog = d_dialog + '(' + ps[i-1] + ')' + d

print(d_dialog)
        
    print(d, p)
#%% 
sc.reformat.header.header_format