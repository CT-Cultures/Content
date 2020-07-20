# -*- coding: utf-8 -*-
"""
Created on Mon Feb 17 00:31:27 2020

@author: VXhpUS
"""


import numpy as np
import pandas as pd
import re
from docx import Document
import time

#%% Import Docx
dir_path = "SCREENPLAYS_PRIVATE"
fnsc = '电影剧本_往来有玉面.docx'
document = Document(dir_path + '/' + fnsc)
num_p = len(document.paragraphs)
script = []
for num in range(num_p):
    script.append(document.paragraphs[num].text)
script = pd.Series(script).rename('pcontent')
script = script[script != ''] # remove empty lines
script = pd.DataFrame(script) # convert to DataFrame
script['ptype'] = np.nan
script['translated'] = np.nan
script['scene_number'] = np.nan
script['content_scene_number'] = np.nan
script['bdcontent'] = script['pcontent'].astype('str').apply(lambda x: x.lstrip(' '))

#%% Function Baidu Translation
import hashlib
from random import randint
import http.client
import urllib
import json

def translate_Baidu(str_to_trans='apple', lang_from='zh', lang_to='en'):
    time.sleep(2)
    #path_translink = 'http://api.fanyi.baidu.com/api/trans/vip/translate'
    path_translink = 'https://fanyi-api.baidu.com/api/trans/vip/translate'
    httpClient = None
    
    appid = '20200218000385369'
    passcode = 'phecLsPI8EnkvjRHQ8HU'
    salt = randint(1e9, 9e9)
    q = str_to_trans
    for_sign = appid + q + str(salt) + passcode
    sign = hashlib.md5(for_sign.encode()).hexdigest()
    
    link_query = (path_translink + '?'
                  + 'appid=' + appid +'&'
                  + 'q=' + urllib.parse.quote(q) + '&'
                  + 'from=' + lang_from + '&'
                  + 'to=' + lang_to + '&'
                  + 'salt=' + str(salt) + '&'
                  + 'sign=' + sign)
    
    try:
        httpClient = http.client.HTTPConnection('api.fanyi.baidu.com')
        httpClient.request('GET', link_query)

        # response是HTTPResponse对象
        response = httpClient.getresponse()
        result_all = response.read().decode("utf-8")
        result = json.loads(result_all)
        print(result['trans_result'][0]['dst'])
        return result['trans_result'][0]['dst']  # str

    except Exception as e:
        print(e)
    finally:
        if httpClient:
           httpClient.close() 

#%% Functions
def clean_line(str_txt):
    re.sub(' ', '', str_txt)
    return str_txt

def extract_elements(script):
    pattern_shot = '.*【.*】'
    pattern_heading = u'.*[0-9]*.【.*】'
    pattern_dialog = u'.*[：:]'
    pattern_text = '.*\[.*\]'
    script.loc[script['bdcontent'].str.match(pattern_shot), 'ptype'] = 'shot'
    script.loc[script['bdcontent'].str.match(pattern_heading), 'ptype'] = 'h'
    script.loc[script['bdcontent'].str.match(pattern_dialog), 'ptype'] = 'd'
    script.loc[script['bdcontent'].str.match(pattern_text), 'ptype'] = 'text'
    return script

def assign_scene_number(script):
    index_headings = script[script['ptype'] == 'h'].index
    for i in range(len(index_headings)-1):
        script.loc[index_headings[i]:index_headings[i+1], 'scene_number'] = i+1
    script.loc[index_headings[i+1]:, 'scene_number'] = i + 1
    return script

def extract_scene_number(script): # extracts scene number from currrent scene heading
    script['content_scene_number'] = np.nan
    script.loc[script['ptype']=='h', 'content_scene_number'] = \
        script.loc[script['ptype']=='h', 'bdcontent'].apply(lambda x: dict(x)['scene_number'])
    script['content_scene_number'].interpolate(method='pad', inplace=True)
    return script

def breakdown_heading(str_heading):
    splitted = re.split('日外到昏外|下午日|午后|黄昏|傍晚|晚上|凌晨|[\s【】。日夜晨昏晚内外]', str_heading)
    scene_number = splitted[0]
    scene_int_ext = None
    scene_time_of_day = None
    splitted = [i for i in splitted[1:] if i  != '']
    if (re.findall('[内外]', str_heading)):
        scene_int_ext = re.findall('[内外]', str_heading)[0]
    if (re.findall('日外到昏外|下午日|午后|黄昏|傍晚|晚上|凌晨|[日夜晨昏晚]', str_heading)):
        scene_time_of_day = re.findall('日外到昏外|下午日|午后|黄昏|傍晚|晚上|凌晨|[日夜晨昏晚]', str_heading)[0]
    scene_name = splitted[0]
    return {'scene_number': scene_number, 'int_ext': scene_int_ext, 
            'scene_name': scene_name, 'time_of_day': scene_time_of_day}

def translate_heading(arow_heading):
    time_of_day = {
            '日外到昏外': 'Day to Evening',
            '下午日': 'Afternoon',
            '午后': 'Afternoon',
            '黄昏': 'Evening',
            '傍晚': 'Late Afternoon',
            '晚上': 'Night',
            '凌晨': 'Late Night',
            '日': 'Day',
            '夜': 'Night',
            '晨': 'Morning',
            '昏': 'Evening',
            '晚': 'Night',
            'None': 'None'
        }
    int_ext = {
            '内': 'INT',
            '外': 'EXT',
            'None': 'None'
        }
    arow_heading = dict(arow_heading)
    li = {
         'scene_number': arow_heading['scene_number'],
         'int_ext': int_ext[str(arow_heading['int_ext'])],
         'scene_name': translate_Baidu(str(arow_heading['scene_name'])),
         'time_of_day': time_of_day[str(arow_heading['time_of_day'])]
         }
    print(li['scene_name'])
    return li
#test =  {'scene_number': '7', 'int_ext': '内', 'scene_name': '北一出站口电子显示牌前', 'time_of_day': '日'}
#translate_heading(test)


def breakdown_dialog(str_dialog): 
    character = re.split('[：:]', str_dialog)[0]
    str_dialog = re.split('[：:]', str_dialog)[1]

    match = re.search(u'（.*）|(.*)', character)
    if (match):
        dialog_wrapped = [(''.join([character[0:match.start()], character[match.end():]]), 'c')]
        dialog_wrapped.append((match.group().lstrip('（').rstrip('）').lstrip('(').rstrip(')'), 'p'))
    else:
        dialog_wrapped = [(character, 'c')]
    
    match = re.search(u'（.*）|(.*)', str_dialog)
    while(match):
        dialog_wrapped.append((str_dialog[0:match.start()], 'd'))
        dialog_wrapped.append((match.group().lstrip('（').rstrip('）').lstrip('(').rstrip(')'), 'p'))
        str_dialog= str_dialog[match.end():]
        match = re.search(u'（.*）|(.*)', str_dialog)
    dialog_wrapped.append((str_dialog, 'd'))
    return dialog_wrapped

def translate_dialog(list_dialog):
    li = []
    for i in range(len(list_dialog)):
        if ((list_dialog[i][1] == 'd') or (list_dialog[i][1] == 'p')):
            li.append((translate_Baidu(list_dialog[i][0]), 'd'))
            print(li)
        else:
            li.append(list_dialog[i])
    return li

      
#%% Operate     

sp = script.copy()
# Manual Cleaning/Editing
sp.loc[417, 'bdcontent'] = sp.loc[417, 'bdcontent'] + '】'     # 417
sp.loc[1145, 'bdcontent'] = sp.loc[1145, 'bdcontent'] + '】'    # 1145       
# 1197

sp = extract_elements(sp)
sp = assign_scene_number(sp) # current scene number

# heading
sp.loc[sp['ptype'] == 'h', 'bdcontent'] = sp.loc[sp['ptype'] == 'h', 'bdcontent'].apply(breakdown_heading)
sp = extract_scene_number(sp) # scene number from import
sp.loc[sp['ptype'] == 'h', 'translated'] = sp.loc[sp['ptype']=='h', 'bdcontent'].apply(translate_heading)

# Action
sp['ptype'].fillna(value='a', inplace=True)
sp.loc[sp['ptype']=='a', 'bdcontent'] = sp.loc[sp['ptype']=='a', 'bdcontent'].apply(lambda x: str(x).lstrip('（').rstrip('）'))
sp.loc[sp['ptype'] == 'a', 'translated'] = sp.loc[sp['ptype'] == 'a', 'bdcontent'].apply(translate_Baidu)

# Dialog
sp.loc[sp['ptype']=='d', 'bdcontent'] = sp.loc[sp['ptype']=='d', 'bdcontent'].apply(breakdown_dialog)
sp.loc[sp['ptype'] == 'd', 'translated'] = sp.loc[sp['ptype']=='d', 'bdcontent'].apply(translate_dialog)


sp.to_csv('./screenplay/spsave.csv')
#%%
from xml.dom import minidom
import lxml.etree as et
dict_basestylenames =  {
    'text':'Normal Text',
    'h': 'Scene Heading',
    'a': 'Action',
    'c': 'Character',
    'p': 'Parenthetical'
    'd': 'Dialogue',
    't': 'Transition',
    's': 'Shot'
    'uc': 'Unspoken Character'
}
        
#%%
def generate_heading_xml(list_headingcontent):
    row = dict(list_headingcontent)
    heading_txt = str(row['int_ext']) + '. ' +  str(row['scene_name']) + ' - ' + str(row['time_of_day'])
    xml = ['    <para>']
    xml.append('      <style basestylename="Scene Heading"/>')
    xml.append('      <text>{}</text>'.format(heading_txt))
    xml.append('    </para>')
    return xml
generate_heading_xml(sp.loc[49, 'bdcontent'])
    
def generate_action_xml(str_actioncontent):
    action_txt = str_actioncontent
    xml = ['    <para>']
    xml.append('      <style basestylename="Action"/>')
    xml.append('      <text>{}</text>'.format(action_txt))
    xml.append('   </para>')
    return xml
#generate_action_xml(sp.loc[53, 'bdcontent'])
    
def generate_dialog_xml(list_dialogcontent):
    xml = []
    for item in list_dialogcontent:
        if (item[1] == 'c'):
            xml.append('    <para>')
            xml.append('      <style basestylename="{}"/>'.format('Character'))
            xml.append('      <text>{}</text>'.format(item[0]))
            xml.append('    </para>')
        elif (item[1] == 'p'):
            xml.append('    <para>')
            xml.append('      <style basestylename="{}"/>'.format('Parenthetical'))
            xml.append('      <text>({})</text>'.format(item[0]))
            xml.append('    </para>')
        elif (item[1] == 'd'):
            xml.append('    <para>')
            xml.append('      <style basestylename="{}"/>'.format('Dialogue'))
            xml.append('      <text>{}</text>'.format(item[0]))
            xml.append('    </para>')
    return xml
#generate_dialog_xml(sp.loc[43, 'bdcontent']).append( generate_action_xml(sp.loc[53, 'pcontent']))

def writing_list_int_ext(rows):
    xml = []
    for row in rows:
        xml.append('        <scene_intro name="{}."/>'.format(dict(row)['int_ext']))
    xml = list(set(xml))
    xml.insert(0, '      <scene_intros>')
    xml.append('      </scene_intros>')
    return xml
#writing_list_int_ext(sp.loc[sp.ptype=='h', 'bdcontent'])

def writing_list_time_of_day(rows):
    xml=[]
    for row in rows:
        xml.append('        <scene_time name="{}"/>'.format(dict(row)['time_of_day']))
    xml = list(set(xml))
    xml.insert(0, '      <scene_times>')
    xml.append('      </scene_times>')
    return xml
#writing_list_time_of_day(sp.loc[sp.ptype=='h', 'bdcontent'])

#%%
## Manual Editing
characters = sp.loc[sp.ptype=='d', 'bdcontent'].apply(lambda x: x[0][0])
character_dict_key = list(set(characters))
character_dict_value = [''] * len(character_dict_key)
dict_characters = dict(zip(character_dict_key, character_dict_value))


#%% Export to XML

xml = ['<?xml version="1.0" encoding="UTF-8"?>']
xml.append('<document type="Open Screenplay Format document" version="30">')
xml.append('  <paragraphs>')
for index, row in sp[sp['scene_number'] >=1].iterrows():
    if (row['ptype'] == 'h'):
        xml = xml + generate_heading_xml(row['bdcontent'])
    elif (row['ptype'] == 'a'):
        xml = xml + generate_action_xml(row['bdcontent'])
    elif (row['ptype'] == 'd'):
        xml = xml + generate_dialog_xml(row['bdcontent'])   
xml.append('  </paragraphs>')
xml.append('  <list>')
xml = xml + writing_list_int_ext(sp.loc[sp.ptype=='h', 'bdcontent'])
xml = xml + writing_list_time_of_day(sp.loc[sp.ptype=='h', 'bdcontent'])
xml.append('  </list>')
xml.append('</document>')

xml_out = '\n'.join(xml)

with open('./screenplay/Leaf_cn_test1.xml', "w", encoding='utf8') as f: 
    f.write(xml_out) 
#%%
xml = ['<?xml version="1.0" encoding="UTF-8"?>']
xml.append('<document type="Open Screenplay Format document" version="30">')
xml.append('  <paragraphs>')
for index, row in sp[sp['scene_number'] >=1].iterrows():
    if (row['ptype'] == 'h'):
        xml = xml + generate_heading_xml(row['translated'])
    elif (row['ptype'] == 'a'):
        xml = xml + generate_action_xml(row['translated'])
    elif (row['ptype'] == 'd'):
        xml = xml + generate_dialog_xml(row['translated'])   
xml.append('  </paragraphs>')
xml.append('  <list>')
xml = xml + writing_list_int_ext(sp.loc[sp.ptype=='h', 'translated'])
xml = xml + writing_list_time_of_day(sp.loc[sp.ptype=='h', 'translated'])
xml.append('  </list>')
xml.append('</document>')

xml_out = '\n'.join(xml)

with open('./screenplay/Leaf_trans_en_test1.xml', "w", encoding='utf8') as f: 
    f.write(xml_out) 

